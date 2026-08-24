"""Meeting Analysis API - 会议内容智能分析

提供文件文本提取和音频转录功能，前端基于提取的文本调用 LLM 完成
客户/商机自动创建及报价单/SOW/WBS 生成。
"""
from __future__ import annotations

import io
import re
import os
import tempfile
from datetime import datetime
from typing import Any, Dict, Optional
from uuid import UUID

import httpx
from fastapi import APIRouter, Depends, File, UploadFile, HTTPException, status, Form, Query
from pydantic import BaseModel, Field
from sqlalchemy import select, desc, func
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.deps import get_current_user_id
from app.core.database import get_db
from app.models.models import MeetingImport, User

router = APIRouter(prefix="/meeting-analysis", tags=["Meeting Analysis"])

# ---------------------------------------------------------------------------
# 文件文本提取
# ---------------------------------------------------------------------------

ALLOWED_TEXT_EXTS = {".txt", ".md", ".csv", ".json", ".log"}
ALLOWED_DOCX_EXTS = {".docx"}
ALLOWED_AUDIO_EXTS = {".mp3", ".wav", ".m4a", ".flac", ".ogg", ".webm", ".mp4"}


def _get_ext(filename: str) -> str:
    return os.path.splitext(filename)[1].lower()


async def _extract_from_text(file_bytes: bytes) -> str:
    """从纯文本文件中提取文本。"""
    # 尝试 UTF-8，失败则尝试 GBK
    for encoding in ("utf-8", "gbk", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except (UnicodeDecodeError, ValueError):
            continue
    return file_bytes.decode("utf-8", errors="replace")


async def _extract_from_docx(file_bytes: bytes) -> str:
    """从 .docx 文件中提取段落文本。"""
    try:
        from docx import Document
    except ImportError:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="服务器未安装 python-docx，无法解析 .docx 文件",
        )
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # 也提取表格内容
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))
    return "\n".join(paragraphs)


@router.post("/extract-text")
async def extract_text(
    file: UploadFile = File(...),
    user_id: str = Depends(get_current_user_id),
):
    """从上传的文件中提取文本内容。

    支持 .txt / .md / .csv / .json / .log / .docx 格式。
    """
    ext = _get_ext(file.filename or "")
    file_bytes = await file.read()

    if not file_bytes:
        raise HTTPException(status_code=400, detail="文件为空")

    if ext in ALLOWED_TEXT_EXTS:
        text = await _extract_from_text(file_bytes)
    elif ext in ALLOWED_DOCX_EXTS:
        text = await _extract_from_docx(file_bytes)
    else:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的文件格式: {ext}。支持的格式: {', '.join(ALLOWED_TEXT_EXTS | ALLOWED_DOCX_EXTS)}",
        )

    return {"text": text, "filename": file.filename, "char_count": len(text)}


# ---------------------------------------------------------------------------
# 音频转录
# ---------------------------------------------------------------------------

def _normalize_base_url(base_url: str) -> str:
    """归一化 base_url，确保以 /v1 结尾。"""
    url = base_url.strip().rstrip("/")
    if re.search(r"/v\d+$", url):
        return url
    return f"{url}/v1"


@router.post("/transcribe")
async def transcribe_audio(
    file: UploadFile = File(...),
    api_key: str = Form(...),
    base_url: str = Form(...),
    model: str = Form(default="whisper-1"),
    user_id: str = Depends(get_current_user_id),
):
    """音频转录 - 调用 OpenAI 兼容的 /v1/audio/transcriptions 接口。

    前端传入 LLM 配置（api_key / base_url / model），
    后端代理调用音频转录 API。
    """
    ext = _get_ext(file.filename or "")
    if ext not in ALLOWED_AUDIO_EXTS:
        raise HTTPException(
            status_code=400,
            detail=f"不支持的音频格式: {ext}。支持的格式: {', '.join(ALLOWED_AUDIO_EXTS)}",
        )

    file_bytes = await file.read()
    if not file_bytes:
        raise HTTPException(status_code=400, detail="音频文件为空")

    # 限制文件大小 25MB（OpenAI Whisper API 限制）
    max_size = 25 * 1024 * 1024
    if len(file_bytes) > max_size:
        raise HTTPException(
            status_code=413,
            detail=f"音频文件过大 ({len(file_bytes) // 1024 // 1024}MB)，最大支持 25MB",
        )

    normalized_url = _normalize_base_url(base_url)
    url = f"{normalized_url}/audio/transcriptions"

    # 写入临时文件供上传
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        async with httpx.AsyncClient(timeout=120.0) as client:
            with open(tmp_path, "rb") as f:
                files = {"file": (file.filename, f, "application/octet-stream")}
                data = {"model": model}
                headers = {"Authorization": f"Bearer {api_key}"}
                resp = await client.post(url, files=files, data=data, headers=headers)

        if resp.status_code != 200:
            err_text = resp.text or "未知错误"
            hints = {
                401: "API Key 无效或已过期",
                402: "账户余额不足",
                403: "无权限使用该模型，或该 provider 不支持音频转录",
                404: "音频转录接口不存在，请确认 provider 是否支持 /v1/audio/transcriptions",
                429: "请求频率超限，请稍后重试",
            }
            hint = hints.get(resp.status_code, "")
            raise HTTPException(
                status_code=status.HTTP_502_BAD_GATEWAY,
                detail=f"音频转录失败 ({resp.status_code}){': ' + hint if hint else ''} | {err_text[:500]}",
            )

        result = resp.json()
        transcript = result.get("text", "")

        return {
            "transcript": transcript,
            "model": model,
            "filename": file.filename,
            "char_count": len(transcript),
        }

    except httpx.TimeoutException:
        raise HTTPException(
            status_code=status.HTTP_504_GATEWAY_TIMEOUT,
            detail="音频转录请求超时（120s），文件可能过大或服务繁忙",
        )
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_502_BAD_GATEWAY,
            detail=f"音频转录调用失败: {str(e)}",
        )
    finally:
        if tmp_path and os.path.exists(tmp_path):
            os.unlink(tmp_path)


# ---------------------------------------------------------------------------
# 会议导入记录 API
# ---------------------------------------------------------------------------

class MeetingImportRecord(BaseModel):
    """前端提交会议导入记录的请求体。"""
    customer_id: str | None = None
    customer_name: str | None = None
    opportunity_id: str | None = None
    opportunity_name: str | None = None
    project_type: str | None = None
    input_type: str = "text"
    input_filename: str | None = None
    meeting_summary: str | None = None
    status: str = "completed"
    has_quotation: bool = False
    has_sow: bool = False
    has_wbs: bool = False


@router.post("/record")
async def record_meeting_import(
    record: MeetingImportRecord,
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    """记录一次会议导入操作。"""
    # 查用户名
    user_result = await db.execute(select(User.username).where(User.id == UUID(user_id)))
    username = user_result.scalar_one_or_none()

    meeting = MeetingImport(
        user_id=UUID(user_id),
        username=username,
        customer_id=UUID(record.customer_id) if record.customer_id else None,
        customer_name=record.customer_name,
        opportunity_id=UUID(record.opportunity_id) if record.opportunity_id else None,
        opportunity_name=record.opportunity_name,
        project_type=record.project_type,
        input_type=record.input_type,
        input_filename=record.input_filename,
        meeting_summary=(record.meeting_summary or "")[:500],
        status=record.status,
        has_quotation=record.has_quotation,
        has_sow=record.has_sow,
        has_wbs=record.has_wbs,
    )
    db.add(meeting)
    await db.commit()
    await db.refresh(meeting)

    return {
        "id": str(meeting.id),
        "created_at": meeting.created_at.isoformat() if meeting.created_at else None,
        "message": "记录已保存",
    }


@router.get("/history")
async def get_meeting_history(
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    user_id: str = Depends(get_current_user_id),
    db: AsyncSession = Depends(get_db),
):
    """获取会议导入历史记录。"""
    # 查总数
    count_result = await db.execute(select(func.count(MeetingImport.id)))
    total = count_result.scalar() or 0

    # 分页查询
    offset = (page - 1) * page_size
    result = await db.execute(
        select(MeetingImport)
        .order_by(desc(MeetingImport.created_at))
        .offset(offset)
        .limit(page_size)
    )
    records = result.scalars().all()

    return {
        "items": [
            {
                "id": str(r.id),
                "user_id": str(r.user_id) if r.user_id else None,
                "username": r.username,
                "customer_id": str(r.customer_id) if r.customer_id else None,
                "customer_name": r.customer_name,
                "opportunity_id": str(r.opportunity_id) if r.opportunity_id else None,
                "opportunity_name": r.opportunity_name,
                "project_type": r.project_type,
                "input_type": r.input_type,
                "input_filename": r.input_filename,
                "meeting_summary": r.meeting_summary,
                "status": r.status,
                "has_quotation": r.has_quotation,
                "has_sow": r.has_sow,
                "has_wbs": r.has_wbs,
                "created_at": r.created_at.isoformat() if r.created_at else None,
            }
            for r in records
        ],
        "total": total,
        "page": page,
        "page_size": page_size,
    }
