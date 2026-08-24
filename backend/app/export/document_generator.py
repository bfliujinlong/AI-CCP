"""AICC Skill 输出文档生成器（PDF / DOCX）。

设计目标：
1. 把 Skill 结构化输出（JSON）转成可下载的人类可读文档
2. 支持 PDF（reportlab）与 DOCX（python-docx）两种格式
3. 通用化处理 —— 不绑定具体 Skill，按 JSON 结构自适应渲染

文档渲染策略：
- 顶层 dict → 章节列表（key 作为章节标题）
- list 元素 → 表格或 bullet 列表
- 嵌套 dict → 子章节
- 标量 → 段落文本
"""

from __future__ import annotations
import io
import os
import json
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Tuple


# 文档存储目录（容器内 / 本地均可）
EXPORT_DIR = Path(os.environ.get("AICC_EXPORT_DIR", "/tmp/aicc_exports") if os.name != "nt" else os.path.join(os.environ.get("TEMP", "C:/temp"), "aicc_exports"))


def ensure_export_dir() -> Path:
    """确保导出目录存在。"""
    EXPORT_DIR.mkdir(parents=True, exist_ok=True)
    return EXPORT_DIR


# ============================================================
# DOCX 生成（python-docx）
# ============================================================

def generate_docx(
    outputs: Dict[str, Any],
    skill_name: str,
    title: str | None = None,
    project_meta: Dict[str, Any] | None = None,
) -> bytes:
    """生成 DOCX 文档（bytes）。

    Args:
        outputs: Skill 执行的结构化输出（dict）
        skill_name: Skill 名称（如 Generate-Quotation）
        title: 文档标题（默认使用 Skill 名）
        project_meta: 项目元数据（客户/项目/日期等），可写进文档头部

    Returns:
        DOCX 文件二进制内容
    """
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # 默认字体
    style = doc.styles["Normal"]
    style.font.name = "Microsoft YaHei"
    style.font.size = Pt(10)

    # 标题
    doc_title = title or skill_name
    h = doc.add_heading(doc_title, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 项目元数据
    if project_meta:
        meta_para = doc.add_paragraph()
        meta_run = meta_para.add_run(" | ".join(f"{k}: {v}" for k, v in project_meta.items() if v))
        meta_run.font.size = Pt(9)
        meta_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # 生成时间
    gen_para = doc.add_paragraph()
    gen_run = gen_para.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | AICC AI")
    gen_run.font.size = Pt(8)
    gen_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.add_paragraph()  # 空行

    # 递归渲染 JSON
    _render_dict_to_docx(doc, outputs)

    # 保存到内存
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def _render_dict_to_docx(doc, obj: Any, level: int = 1) -> None:
    """递归把 dict/list/标量 渲染到 DOCX。"""
    from docx.shared import RGBColor

    if isinstance(obj, dict):
        for key, value in obj.items():
            display_key = _humanize_key(key)
            if isinstance(value, (dict, list)) and value:
                doc.add_heading(display_key, level=min(level + 1, 4))
                _render_dict_to_docx(doc, value, level + 1)
            else:
                p = doc.add_paragraph()
                run = p.add_run(f"{display_key}: ")
                run.bold = True
                p.add_run(str(value) if value is not None else "-")
    elif isinstance(obj, list):
        if not obj:
            doc.add_paragraph("(无)")
            return
        # 如果 list 元素是 dict 且字段一致 → 表格
        if isinstance(obj[0], dict) and all(isinstance(x, dict) for x in obj):
            _render_list_of_dict_as_table(doc, obj)
        else:
            # bullet 列表
            for item in obj:
                p = doc.add_paragraph(str(item), style="List Bullet")
    else:
        doc.add_paragraph(str(obj))


def _render_list_of_dict_as_table(doc, items: List[Dict[str, Any]]) -> None:
    """把 dict 列表渲染为 DOCX 表格。"""
    if not items:
        return
    keys = list(items[0].keys())
    # 列数 ≤ 6，超出则截断展示前 6 列
    if len(keys) > 6:
        keys = keys[:6]
    table = doc.add_table(rows=1 + len(items), cols=len(keys))
    table.style = "Light Grid Accent 1"
    # 表头
    hdr_cells = table.rows[0].cells
    for i, k in enumerate(keys):
        hdr_cells[i].text = _humanize_key(k)
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    # 数据行
    for r, item in enumerate(items, start=1):
        for c, k in enumerate(keys):
            v = item.get(k)
            if isinstance(v, (dict, list)):
                cell_text = json.dumps(v, ensure_ascii=False)[:80]
            else:
                cell_text = str(v) if v is not None else ""
            table.rows[r].cells[c].text = cell_text


# ============================================================
# PDF 生成（reportlab）
# ============================================================

def generate_pdf(
    outputs: Dict[str, Any],
    skill_name: str,
    title: str | None = None,
    project_meta: Dict[str, Any] | None = None,
) -> bytes:
    """生成 PDF 文档（bytes）。"""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
    )
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # 注册中文字体（尽量找系统已有的中文字体）
    _register_chinese_font()

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2 * cm, rightMargin=2 * cm,
        topMargin=2 * cm, bottomMargin=2 * cm,
        title=title or skill_name,
    )
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1", parent=styles["Heading1"], fontSize=18, spaceAfter=12, alignment=1)
    h2 = ParagraphStyle("h2", parent=styles["Heading2"], fontSize=14, spaceBefore=10, spaceAfter=6)
    h3 = ParagraphStyle("h3", parent=styles["Heading3"], fontSize=12, spaceBefore=8, spaceAfter=4)
    body = ParagraphStyle("body", parent=styles["BodyText"], fontSize=10, leading=14)
    meta_st = ParagraphStyle("meta", parent=styles["BodyText"], fontSize=8, textColor=colors.grey, alignment=1)

    story = []
    doc_title = title or skill_name
    story.append(Paragraph(doc_title, h1))
    story.append(Paragraph(f"Skill: {skill_name}", meta_st))

    if project_meta:
        meta_str = " | ".join(f"{k}: {v}" for k, v in project_meta.items() if v)
        if meta_str:
            story.append(Paragraph(meta_str, meta_st))

    story.append(Paragraph(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", meta_st))
    story.append(Spacer(1, 0.5 * cm))

    _render_dict_to_pdf(story, outputs, h2, h3, body)

    doc.build(story)
    buf.seek(0)
    return buf.getvalue()


def _render_dict_to_pdf(story, obj: Any, h2, h3, body) -> None:
    """递归把 dict/list 渲染到 PDF story。"""
    from reportlab.platypus import Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors

    if isinstance(obj, dict):
        for key, value in obj.items():
            display_key = _humanize_key(key)
            if isinstance(value, (dict, list)) and value:
                story.append(Paragraph(display_key, h2))
                _render_dict_to_pdf(story, value, h2, h3, body)
            else:
                story.append(Paragraph(f"<b>{display_key}:</b> {_escape(str(value) if value is not None else '-')}", body))
    elif isinstance(obj, list):
        if not obj:
            story.append(Paragraph("(无)", body))
            return
        if isinstance(obj[0], dict) and all(isinstance(x, dict) for x in obj):
            _render_list_of_dict_as_pdf_table(story, obj)
        else:
            for item in obj:
                story.append(Paragraph(f"• {_escape(str(item))}", body))
    else:
        story.append(Paragraph(_escape(str(obj)), body))


def _render_list_of_dict_as_pdf_table(story, items) -> None:
    """dict 列表 → PDF 表格。"""
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib import colors

    keys = list(items[0].keys())
    if len(keys) > 6:
        keys = keys[:6]
    data = [[_humanize_key(k) for k in keys]]
    for item in items:
        row = []
        for k in keys:
            v = item.get(k)
            if isinstance(v, (dict, list)):
                cell_text = json.dumps(v, ensure_ascii=False)[:80]
            else:
                cell_text = str(v) if v is not None else ""
            row.append(_escape(cell_text))
        data.append(row)

    table = Table(data, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#534AB7")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "ChineseFont"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 6),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.grey),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(table)


def _register_chinese_font() -> bool:
    """尝试注册中文字体；失败则用 Helvetica（中文会显示为方块，但用户主要看 DOCX）。"""
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    font_paths = [
        # Windows
        "C:/Windows/Fonts/msyh.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/simsun.ttc",
        # Linux/Mac 常见路径
        "/usr/share/fonts/truetype/wqy/wqy-zenhei.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
    ]
    for fp in font_paths:
        if Path(fp).exists():
            try:
                pdfmetrics.registerFont(TTFont("ChineseFont", fp))
                return True
            except Exception:
                continue
    return False


def _escape(text: str) -> str:
    """转义 ReportLab XML 特殊字符。"""
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def _humanize_key(key: str) -> str:
    """把 camelCase / snake_case 转中文友好标题。"""
    # 常见 key 中文映射
    mapping = {
        "summary": "摘要",
        "title": "标题",
        "name": "名称",
        "description": "描述",
        "questions": "问题清单",
        "decisions": "决策项",
        "action_items": "行动项",
        "risks": "风险点",
        "recommendations": "建议",
        "phases": "阶段",
        "tasks": "任务",
        "total_cost": "总成本",
        "total_days": "总工期",
        "total_person_days": "总人天",
        "currency": "货币",
        "risk_premium": "风险溢价",
        "categories": "分类",
        "checklist": "检查清单",
        "rollback_plan": "回滚方案",
        "deliverables": "交付物",
        "assumptions": "假设与前提",
        "milestones": "里程碑",
        "chapters": "章节",
        "waves": "切换波次",
        "next_agenda": "下次会议议程",
        "controls": "控制项",
        "compliance_score": "合规得分",
        "critical_gaps": "关键差距",
        "optimizations": "优化项",
        "short_term_actions": "短期措施",
        "long_term_actions": "中长期措施",
        "estimated_savings": "预估节省",
        "roi": "投资回报率",
    }
    if key in mapping:
        return mapping[key]
    # snake_case → Title Case
    return key.replace("_", " ").title()


# ============================================================
# 主入口
# ============================================================

def generate_document(
    outputs: Dict[str, Any],
    skill_name: str,
    format: str,
    title: str | None = None,
    project_meta: Dict[str, Any] | None = None,
) -> Tuple[bytes, str]:
    """统一入口。

    Args:
        outputs: Skill 输出 JSON
        skill_name: Skill 名
        format: "docx" / "pdf" / "md" / "none"
        title: 文档标题
        project_meta: 头部元数据

    Returns:
        (文件二进制, 文件扩展名)
    """
    fmt = (format or "docx").lower()
    if fmt == "docx":
        content = generate_docx(outputs, skill_name, title=title, project_meta=project_meta)
        return content, "docx"
    elif fmt == "pdf":
        content = generate_pdf(outputs, skill_name, title=title, project_meta=project_meta)
        return content, "pdf"
    elif fmt == "md":
        content = _generate_markdown(outputs, skill_name, title=title, project_meta=project_meta).encode("utf-8")
        return content, "md"
    else:
        # 不生成文档
        return b"", ""


def _generate_markdown(outputs: Dict[str, Any], skill_name: str, title: str | None = None, project_meta: Dict[str, Any] | None = None) -> str:
    """生成 Markdown。"""
    lines = [f"# {title or skill_name}", f"**Skill**: {skill_name}"]
    if project_meta:
        lines.append(" | ".join(f"**{k}**: {v}" for k, v in project_meta.items() if v))
    lines.append(f"**生成时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    lines.append("")
    _render_dict_to_md(lines, outputs)
    return "\n".join(lines)


def _render_dict_to_md(lines: List[str], obj: Any, level: int = 2) -> None:
    if isinstance(obj, dict):
        for key, value in obj.items():
            display_key = _humanize_key(key)
            if isinstance(value, (dict, list)) and value:
                lines.append(f"\n{'#' * level} {display_key}\n")
                _render_dict_to_md(lines, value, min(level + 1, 6))
            else:
                lines.append(f"- **{display_key}**: {value if value is not None else '-'}")
    elif isinstance(obj, list):
        if not obj:
            lines.append("- (无)")
            return
        if isinstance(obj[0], dict) and all(isinstance(x, dict) for x in obj):
            # 渲染为 markdown 表格
            keys = list(obj[0].keys())[:6]
            lines.append("| " + " | ".join(_humanize_key(k) for k in keys) + " |")
            lines.append("| " + " | ".join("---" for _ in keys) + " |")
            for item in obj:
                row = []
                for k in keys:
                    v = item.get(k)
                    if isinstance(v, (dict, list)):
                        cell_text = json.dumps(v, ensure_ascii=False)[:80]
                    else:
                        cell_text = str(v) if v is not None else ""
                    row.append(cell_text.replace("|", "\\|"))
                lines.append("| " + " | ".join(row) + " |")
        else:
            for item in obj:
                lines.append(f"- {item}")
    else:
        lines.append(str(obj))


def save_document(
    outputs: Dict[str, Any],
    skill_name: str,
    format: str,
    execution_id: str | None = None,
    title: str | None = None,
    project_meta: Dict[str, Any] | None = None,
) -> Dict[str, str]:
    """生成文档并保存到磁盘，返回下载信息。

    Returns:
        {"filename": "...", "path": "...", "format": "docx", "url": "/api/v1/skills/documents/..."}
    """
    content, ext = generate_document(outputs, skill_name, format, title=title, project_meta=project_meta)
    if not ext:
        return {}

    export_dir = ensure_export_dir()
    safe_skill = "".join(c if c.isalnum() or c in "-_" else "_" for c in skill_name)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    exec_id = execution_id or timestamp
    filename = f"{safe_skill}_{exec_id}.{ext}"
    filepath = export_dir / filename

    filepath.write_bytes(content)

    return {
        "filename": filename,
        "path": str(filepath),
        "format": ext,
        "size_bytes": str(len(content)),
        "url": f"/api/v1/skills/documents/{filename}",
    }